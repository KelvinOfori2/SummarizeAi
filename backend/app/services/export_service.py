from fastapi import HTTPException
from fastapi.responses import Response
from app.models.summary import Summary
from app.core.logging import get_logger
import io

logger = get_logger(__name__)


def export_as_txt(summary: Summary) -> Response:
    content_lines = [
        f"SUMMARY — {summary.title or 'Untitled'}",
        "=" * 60,
        f"Algorithm     : {summary.algorithm.upper()}",
        f"Original Words: {summary.original_word_count}",
        f"Summary Words : {summary.summary_word_count}",
        f"Compression   : {round((1 - summary.compression_ratio) * 100, 1)}%",
        f"Generated At  : {summary.created_at.strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "SUMMARY",
        "-" * 60,
        summary.summary_text,
        "",
        "ORIGINAL TEXT",
        "-" * 60,
        summary.original_text,
    ]
    content = "\n".join(content_lines)
    return Response(
        content=content.encode("utf-8"),
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="summary_{summary.id[:8]}.txt"'},
    )


def export_as_pdf(summary: Summary) -> Response:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("title", parent=styles["Title"], fontSize=18, spaceAfter=6)
        heading_style = ParagraphStyle("heading", parent=styles["Heading2"], fontSize=13, spaceAfter=4)
        body_style = ParagraphStyle("body", parent=styles["Normal"], fontSize=10, leading=16)
        meta_style = ParagraphStyle("meta", parent=styles["Normal"], fontSize=9,
                                    textColor=colors.grey, spaceAfter=4)

        story = []
        story.append(Paragraph(summary.title or "Summary Report", title_style))
        story.append(Spacer(1, 0.3 * cm))

        meta_data = [
            ["Algorithm", summary.algorithm.upper()],
            ["Original Words", str(summary.original_word_count)],
            ["Summary Words", str(summary.summary_word_count)],
            ["Compression", f"{round((1 - summary.compression_ratio) * 100, 1)}%"],
            ["Processing Time", f"{summary.processing_time}s"],
            ["Generated", summary.created_at.strftime("%Y-%m-%d %H:%M:%S")],
        ]
        table = Table(meta_data, colWidths=[4 * cm, 10 * cm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f0f0")),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.5 * cm))

        story.append(Paragraph("Summary", heading_style))
        for para in summary.summary_text.split("\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), body_style))
                story.append(Spacer(1, 0.2 * cm))

        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph("Original Text", heading_style))
        for para in summary.original_text.split("\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), body_style))
                story.append(Spacer(1, 0.1 * cm))

        doc.build(story)
        buffer.seek(0)

        return Response(
            content=buffer.read(),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="summary_{summary.id[:8]}.pdf"'},
        )
    except Exception as e:
        logger.error(f"PDF export error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate PDF export")


def export_as_docx(summary: Summary) -> Response:
    try:
        import docx
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        document = docx.Document()

        document.add_heading(summary.title or "Summary Report", level=0)

        table = document.add_table(rows=6, cols=2)
        table.style = "Table Grid"
        meta_rows = [
            ("Algorithm", summary.algorithm.upper()),
            ("Original Words", str(summary.original_word_count)),
            ("Summary Words", str(summary.summary_word_count)),
            ("Compression", f"{round((1 - summary.compression_ratio) * 100, 1)}%"),
            ("Processing Time", f"{summary.processing_time}s"),
            ("Generated", summary.created_at.strftime("%Y-%m-%d %H:%M:%S")),
        ]
        for i, (k, v) in enumerate(meta_rows):
            row = table.rows[i]
            row.cells[0].text = k
            row.cells[1].text = v

        document.add_paragraph()
        document.add_heading("Summary", level=2)
        document.add_paragraph(summary.summary_text)

        document.add_paragraph()
        document.add_heading("Original Text", level=2)
        document.add_paragraph(summary.original_text)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        return Response(
            content=buffer.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="summary_{summary.id[:8]}.docx"'},
        )
    except Exception as e:
        logger.error(f"DOCX export error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate DOCX export")
